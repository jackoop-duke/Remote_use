#!/usr/bin/env python3
"""Minimal Markdown -> DOCX converter tailored for the survey report.
Supports: # headings (1-4), paragraphs, bullet/numbered lists, pipe tables,
**bold**, *italic*, `code`, [text](url) links (rendered as text + URL), horizontal rules.
Usage: md2docx.py input.md output.docx
"""
import re, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INLINE = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')

def set_cell_shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_inline(par, text, size=None, bold_all=False):
    parts = INLINE.split(text)
    for p in parts:
        if not p:
            continue
        if p.startswith('**') and p.endswith('**'):
            r = par.add_run(p[2:-2]); r.bold = True
        elif p.startswith('*') and p.endswith('*') and len(p) > 2:
            r = par.add_run(p[1:-1]); r.italic = True
        elif p.startswith('`') and p.endswith('`'):
            r = par.add_run(p[1:-1]); r.font.name = 'Consolas'
        elif p.startswith('['):
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', p)
            r = par.add_run(m.group(1)); r.underline = True; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            if m.group(2) != m.group(1):
                r2 = par.add_run(f' <{m.group(2)}>'); r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        else:
            r = par.add_run(p)
        if bold_all: r.bold = True
        if size: r.font.size = Pt(size)

def flush_table(doc, rows):
    rows = [r for r in rows if not re.match(r'^\s*\|?\s*:?-{2,}', r)]
    cells = [[c.strip() for c in r.strip().strip('|').split('|')] for r in rows]
    if not cells: return
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(cells):
        for j in range(ncol):
            txt = row[j] if j < len(row) else ''
            cell = t.cell(i, j); cell.text = ''
            par = cell.paragraphs[0]
            add_inline(par, txt.replace('<br>', '\n'), size=8, bold_all=(i == 0))
            if i == 0: set_cell_shade(cell, 'D9E2F3')
    doc.add_paragraph()

def main(src, dst):
    doc = Document()
    st = doc.styles['Normal']; st.font.name = 'Yu Gothic'; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    for s in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
        doc.styles[s].font.name = 'Yu Gothic'
        doc.styles[s].element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Cm(2.0); sec.top_margin = sec.bottom_margin = Cm(2.0)
    lines = open(src, encoding='utf-8').read().splitlines()
    table_buf = []; i = 0; first_h1 = True
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('|'):
            table_buf.append(line); i += 1; continue
        if table_buf:
            flush_table(doc, table_buf); table_buf = []
        s = line.rstrip()
        if not s.strip():
            i += 1; continue
        m = re.match(r'^(#{1,4})\s+(.*)', s)
        if m:
            lvl = len(m.group(1)); text = m.group(2).strip()
            if lvl == 1 and first_h1:
                p = doc.add_heading(level=0); add_inline(p, text); first_h1 = False
            else:
                p = doc.add_heading(level=lvl); add_inline(p, text)
            i += 1; continue
        if re.match(r'^\s*-{3,}\s*$', s):
            doc.add_paragraph().add_run('─' * 40).font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF); i += 1; continue
        m = re.match(r'^(\s*)[-*]\s+(.*)', s)
        if m:
            indent = len(m.group(1)) // 2
            p = doc.add_paragraph(style='List Bullet' if indent == 0 else 'List Bullet 2'); add_inline(p, m.group(2)); i += 1; continue
        m = re.match(r'^\s*\d+[.)]\s+(.*)', s)
        if m:
            p = doc.add_paragraph(style='List Number'); add_inline(p, m.group(1)); i += 1; continue
        if s.startswith('>'):
            p = doc.add_paragraph(); add_inline(p, s.lstrip('> ')); p.paragraph_format.left_indent = Cm(1); i += 1; continue
        # paragraph: merge consecutive non-empty plain lines
        buf = [s]
        while i + 1 < len(lines) and lines[i+1].strip() and not re.match(r'^(#{1,4}\s|\s*[-*]\s|\s*\d+[.)]\s|\||>|-{3,})', lines[i+1]):
            i += 1; buf.append(lines[i].rstrip())
        p = doc.add_paragraph(); add_inline(p, ' '.join(buf)); p.paragraph_format.space_after = Pt(6)
        i += 1
    if table_buf: flush_table(doc, table_buf)
    doc.save(dst)

if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2])
